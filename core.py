import os
import unicodedata
import warnings

# Suppress warnings
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["HF_HUB_DISABLE_PROGRESS_BARS"] = "1"
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

import google.generativeai as genai
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document as LangchainDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_community.chat_message_histories import RedisChatMessageHistory
from langchain_core.messages import HumanMessage, AIMessage


class InsuranceLawChatbot:
    def __init__(self, docx_files_path, api_key=None, redis_url="redis://localhost:6379"):
        """
        Initialize the RAG chatbot for Kazakhstani insurance laws using Google Gemini.

        Args:
            docx_files_path: Path to folder containing .docx files or list of file paths
            api_key: Google API key (or set GOOGLE_API_KEY env variable)
            redis_url: Redis connection URL (default: redis://localhost:6379)
        """
        if api_key:
            os.environ["GOOGLE_API_KEY"] = api_key

        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

        self.docx_files_path = docx_files_path
        self.redis_url = redis_url
        self.vectorstore = None
        self.retriever = None
        self.chain = None
        self.model = None

    @staticmethod
    def normalize_text(text):
        """
        Normalize text to avoid tokenizer issues with special characters.
        Removes or replaces problematic characters.
        """
        if not text or not isinstance(text, str):
            return ""

        # Normalize unicode (NFC normalization)
        text = unicodedata.normalize('NFKC', text)

        # Remove or replace characters that cause tokenizer issues
        safe_chars = []
        for char in text:
            code = ord(char)
            # Basic Latin, Cyrillic, common punctuation, spaces
            if (32 <= code <= 126 or  # Basic Latin
                    1024 <= code <= 1279 or  # Cyrillic
                    char in ' .,;:!?-—–()[]{}\"\'`' or  # Punctuation
                    char.isspace()):
                safe_chars.append(char)
            else:
                safe_chars.append(' ')

        result = ''.join(safe_chars)
        # Remove multiple spaces
        result = ' '.join(result.split())

        return result

    def load_docx_files(self):
        """Load all .docx files from the specified path with enhanced metadata."""
        documents = []

        # Handle both single file and directory
        if isinstance(self.docx_files_path, str):
            if os.path.isfile(self.docx_files_path):
                files = [self.docx_files_path]
            else:
                files = [
                    os.path.join(self.docx_files_path, f)
                    for f in os.listdir(self.docx_files_path)
                    if f.endswith('.docx')
                ]
        else:
            files = self.docx_files_path

        for file_path in files:
            doc = Document(file_path)
            full_text = []
            
            filename = os.path.basename(file_path)

            # Extract text from paragraphs with structure preservation
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    # Try to preserve heading/structure info
                    if para.style and 'Heading' in para.style.name:
                        full_text.append(f"\n{'='*50}\n{text}\n{'='*50}")
                    else:
                        full_text.append(text)

            # Extract text from tables with table structure
            for table_idx, table in enumerate(doc.tables):
                full_text.append(f"\n[ТАБЛИЦА {table_idx + 1} из {filename}]")
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        full_text.append(" | ".join(row_data))

            # Create a document with enhanced metadata
            text_content = '\n'.join(full_text)
            if text_content:  # Only add if document has content
                documents.append(
                    LangchainDocument(
                        page_content=text_content,
                        metadata={
                            "source": filename,
                            "type": self._get_doc_type(filename)
                        }
                    )
                )

        print(f"Loaded {len(documents)} documents: {[d.metadata['source'] for d in documents]}")
        return documents

    @staticmethod
    def _get_doc_type(filename):
        """Categorize document type for better context."""
        if "ГКРК" in filename:
            return "Гражданский кодекс"
        elif "ОГПОВТС" in filename:
            return "Обязательное страхование ТС"
        elif "перевозчика" in filename:
            return "Страхование перевозчика"
        elif "туриста" in filename:
            return "Страхование туристов"
        elif "экологии" in filename:
            return "Страхование окружающей среды"
        elif "Нотариусы" in filename:
            return "Страхование нотариусов"
        elif "опасные" in filename:
            return "Страхование опасных объектов"
        else:
            return "Прочие документы"

    def chunk_documents(self, documents):
        """Split documents into smaller chunks for better retrieval."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        chunks = text_splitter.split_documents(documents)
        print(f"Created {len(chunks)} chunks from documents")
        return chunks

    def create_vectorstore(self, chunks):
        """Create a vector database from document chunks using Google embeddings."""
        # Use Google embeddings
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001"
        )

        # Create Chroma vector store
        self.vectorstore = Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory="./insurance_law_db"
        )

        # Create retriever - set to 20 chunks for comprehensive context
        self.retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 10}  # Retrieve top 20 most relevant chunks
        )

        print("Vector store created successfully")
        return self.vectorstore

    def setup_chain(self, model_name="gemini-1.5-flash"):
        """
        Set up the RAG chain with Google Gemini.
        """
        from langchain_google_genai import ChatGoogleGenerativeAI
        self.model = ChatGoogleGenerativeAI(
            model=model_name,
            temperature=0,
            convert_system_message_to_human=True
        )
        # Enhanced prompt to return ALL relevant information
        template = """Вы - эксперт по казахскому страховому законодательству.

    КРИТИЧЕСКИЕ ИНСТРУКЦИИ:
    1. ВЫДАВАЙТЕ ВСЮ РЕЛЕВАНТНУЮ ИНФОРМАЦИЮ из предоставленного контекста
    2. Если документ содержит список (пункты, случаи, примеры) - ПЕРЕЧИСЛИТЕ ВСЕ без пропусков
    3. Не сокращайте информацию - покажите полный список всех пунктов/случаев
    4. Для каждого пункта/случая дайте четкое объяснение
    5. Используйте нумерацию или маркеры для читаемости
    6. КРИТИЧЕСКИ ВАЖНО: Если вопрос касается перечисления - выдайте ВСЕ пункты, не ограничивайте себя

    История разговора:
    {chat_history}

    КОНТЕКСТ ИЗ ДОКУМЕНТОВ (используйте ВСЮ релевантную информацию):
    {context}

    Текущий вопрос: {question}

    ОТВЕТ ДОЛЖЕН:
    - Содержать ВСЕ случаи/пункты/примеры из контекста
    - Быть полным и развернутым при перечислении
    - Указывать точные номера статей и пунктов
    - Использовать нумерацию для ясности
    - Если ответ найден в документах - добавить маркер: [ИСТОЧНИКИ_НАЙДЕНЫ]

    Ответ:"""

        prompt = ChatPromptTemplate.from_template(template)
        # Format documents function
        def format_docs(docs):
            if not docs:
                return "Релевантная информация не найдена."
            docs_by_source = {}
            for doc in docs:
                source = doc.metadata.get("source", "Неизвестный источник")
                if source not in docs_by_source:
                    docs_by_source[source] = []
                docs_by_source[source].append(doc.page_content)
            formatted = []
            for source, contents in docs_by_source.items():
                formatted.append(f"\n📄 Из '{source}':")
                formatted.append("=" * 40)
                for content in contents:
                    formatted.append(content)
                formatted.append("=" * 40)
            return "\n".join(formatted)
        # Format chat history
        def format_chat_history(history):
            if not history:
                return "Нет истории."
            formatted = []
            for msg in history[-6:]:
                if isinstance(msg, HumanMessage):
                    formatted.append(f"Вопрос: {msg.content}")
                elif isinstance(msg, AIMessage):
                    formatted.append(f"Ответ: {msg.content}")
            return "\n".join(formatted)
        self.format_docs = format_docs
        self.format_chat_history = format_chat_history
        self.chain = (
                prompt
                | self.model
                | StrOutputParser()
        )
        print(f"RAG chain ready with {model_name}")
        return self.chain

    def build(self, model_name="gemini-1.5-flash"):
        """Build the complete RAG system."""
        print("Building RAG system with Google Gemini...")

        # Step 1: Load documents
        documents = self.load_docx_files()

        # Step 2: Chunk documents
        chunks = self.chunk_documents(documents)

        # Step 3: Create vector store
        self.create_vectorstore(chunks)

        # Step 4: Setup chain
        self.setup_chain(model_name)

        print("RAG system built successfully!")

    def ask(self, question, user_id="default"):
        """
        Ask a question about insurance laws.

        Args:
            question: Question in Russian or Kazakh
            user_id: Unique identifier for the user

        Returns:
            dict with 'answer', 'sources', and 'has_sources' flag
        """
        if not self.chain:
            raise ValueError("Please run build() first to initialize the system")

        # Normalize text
        normalized_question = self.normalize_text(question)

        # Safety check
        if not normalized_question or len(normalized_question.strip()) < 2:
            return {
                "answer": "Извините, я не смог обработать ваш вопрос. Попробуйте переформулировать.",
                "sources": [],
                "source_documents": [],
                "has_sources": False
            }

        try:
            # Get or create Redis chat history for this user
            chat_history = RedisChatMessageHistory(
                session_id=f"user:{user_id}",
                url=self.redis_url,
                ttl=86400
            )

            # Get relevant documents with error handling
            try:
                relevant_docs = self.retriever.invoke(normalized_question)
            except Exception as e:
                print(f"Retriever error: {e}")
                simple_question = ''.join(c for c in normalized_question if c.isalnum() or c.isspace())
                if simple_question:
                    try:
                        relevant_docs = self.retriever.invoke(simple_question)
                    except Exception as e2:
                        print(f"Simple question retrieval also failed: {e2}")
                        relevant_docs = []
                else:
                    relevant_docs = []

            # Format chat history
            formatted_history = self.format_chat_history(chat_history.messages)

            # Prepare context
            if relevant_docs:
                formatted_context = self.format_docs(relevant_docs)
            else:
                formatted_context = "Релевантные документы не найдены."

            # Build the input for the chain
            chain_input = {
                "context": formatted_context,
                "question": normalized_question,
                "chat_history": formatted_history
            }

            # Get answer from chain
            full_answer = self.chain.invoke(chain_input)

            # Check if LLM decided to include sources marker
            has_sources_marker = "[ИСТОЧНИКИ_НАЙДЕНЫ]" in full_answer
            
            # Remove the marker from the answer for display
            answer = full_answer.replace("[ИСТОЧНИКИ_НАЙДЕНЫ]", "").strip()

            # Add to chat history
            chat_history.add_user_message(question)
            chat_history.add_ai_message(answer)

            return {
                "answer": answer,
                "sources": [doc.metadata["source"] for doc in relevant_docs] if relevant_docs else [],
                "source_documents": relevant_docs,
                "has_sources": has_sources_marker and len(relevant_docs) > 0
            }

        except Exception as e:
            print(f"Error in ask(): {e}")
            import traceback
            traceback.print_exc()
            return {
                "answer": "Извините, произошла ошибка при обработке вашего вопроса. Попробуйте задать вопрос по-другому или используйте команду /clear для начала нового разговора.",
                "sources": [],
                "source_documents": [],
                "has_sources": False
            }

    def clear_history(self, user_id="default"):
        """Clear chat history for a specific user."""
        chat_history = RedisChatMessageHistory(
            session_id=f"user:{user_id}",
            url=self.redis_url
        )
        chat_history.clear()

    def get_history(self, user_id="default"):
        """Get chat history for a specific user."""
        chat_history = RedisChatMessageHistory(
            session_id=f"user:{user_id}",
            url=self.redis_url
        )
        return chat_history.messages

    def load_existing_vectorstore(self):
        """Load previously created vector store without rebuilding."""
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

        self.vectorstore = Chroma(
            persist_directory="./insurance_law_db",
            embedding_function=embeddings
        )

        # Create retriever - set to 20 chunks for comprehensive context
        self.retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 10}  # Retrieve top 20 most relevant chunks
        )

        print("Loaded existing vector store")

    def get_knowledge_base_status(self) -> dict:
        """Get information about loaded knowledge base."""
        from manager import KnowledgeManager
        km = KnowledgeManager()
        return km.get_knowledge_base_info()


# Example usage
if __name__ == "__main__":
    chatbot = InsuranceLawChatbot(
        docx_files_path="./insurance_laws",
        api_key="your-google-api-key"  # Get from https://ai.google.dev
    )

    # Build the RAG system (do this once)
    chatbot.build(model_name="gemini-1.5-flash")

    # For subsequent runs:
    # chatbot.load_existing_vectorstore()
    # chatbot.setup_chain("gemini-1.5-flash")

    # Ask questions
    questions = [
        "Какие виды страхования предусмотрены законом?",
        "Каковы права страхователя?",
        "Какие требования к страховым компаниям?",
        "Что такое обязательное страхование?"
    ]

    for question in questions:
        print(f"\n{'=' * 60}")
        print(f"Вопрос: {question}")
        print('=' * 60)
        result = chatbot.ask(question)
        print(f"Ответ: {result['answer']}")
        print(f"\nИсточники: {', '.join(set(result['sources']))}")
