"""
Knowledge base management system for insurance law assistant.
Handles document loading, updating, and terminology extraction.
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict
import hashlib
from docx import Document
from langchain_core.documents import Document as LangchainDocument


class KnowledgeManager:
    """Manages document knowledge base and terminology extraction."""

    def __init__(self, db_path="./knowledge_base"):
        self.db_path = Path(db_path)
        self.db_path.mkdir(exist_ok=True)
        self.metadata_file = self.db_path / "documents_metadata.json"
        self.terminology_file = self.db_path / "terminology.json"
        self.loaded_docs = {}
        self.terminology = {}
        self._load_metadata()
        self._load_terminology()

    def _load_metadata(self):
        """Load document metadata from JSON file."""
        if self.metadata_file.exists():
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                self.loaded_docs = json.load(f)

    def _save_metadata(self):
        """Save document metadata to JSON file."""
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump(self.loaded_docs, f, ensure_ascii=False, indent=2)

    def _load_terminology(self):
        """Load terminology and domain-specific terms."""
        if self.terminology_file.exists():
            with open(self.terminology_file, 'r', encoding='utf-8') as f:
                self.terminology = json.load(f)
        else:
            self._extract_default_terminology()

    def _save_terminology(self):
        """Save terminology to JSON file."""
        with open(self.terminology_file, 'w', encoding='utf-8') as f:
            json.dump(self.terminology, f, ensure_ascii=False, indent=2)

    def _extract_default_terminology(self):
        """Extract default insurance industry terminology in Kazakhstan."""
        self.terminology = {
            "страхование": "Contract or relationship between insurer and insured",
            "страховщик": "Insurance company providing insurance",
            "страхователь": "Person purchasing insurance",
            "застрахованное лицо": "Person covered by insurance",
            "ОСАГО": "Обязательное страхование гражданской ответственности (Mandatory auto liability insurance)",
            "ГКРК": "Гражданский кодекс Республики Казахстан (Civil Code of RK)",
            "страховой полис": "Insurance contract document",
            "страховая премия": "Insurance fee paid by insured",
            "страховое возмещение": "Insurance payment to insured",
            "убыток": "Loss or damage covered by insurance",
            "страховой случай": "Event covered by insurance contract",
            "франшиза": "Deductible - amount insured pays before insurance",
            "лимит": "Maximum amount insurance will pay",
            "ответственность": "Liability or responsibility",
            "ущерб": "Damage or harm requiring compensation",
        }
        self._save_terminology()

    def get_file_hash(self, file_path: str) -> str:
        """Calculate file hash to detect changes."""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            buf = f.read()
            hasher.update(buf)
        return hasher.hexdigest()

    def has_document_changed(self, file_path: str) -> bool:
        """Check if document has been modified since last load."""
        filename = os.path.basename(file_path)
        if filename not in self.loaded_docs:
            return True

        current_hash = self.get_file_hash(file_path)
        stored_hash = self.loaded_docs[filename].get("hash")
        return current_hash != stored_hash

    def load_docx_files(self, directory: str, extract_terminology: bool = True) -> List[LangchainDocument]:
        """
        Load all DOCX files from directory with enhanced extraction.

        Args:
            directory: Path to directory with DOCX files
            extract_terminology: Whether to extract new terminology from documents
        """
        documents = []
        docx_files = list(Path(directory).glob("*.docx"))

        print(f"📚 Found {len(docx_files)} DOCX files")

        for file_path in docx_files:
            filename = file_path.name

            # Check if document needs updating
            if not self.has_document_changed(str(file_path)):
                print(f"✓ {filename} (not changed, skipping)")
                continue

            print(f"📖 Loading: {filename}")
            doc = Document(file_path)
            full_text = []

            # Extract text with structure
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    if para.style and 'Heading' in para.style.name:
                        full_text.append(f"\n{'=' * 60}\n{text}\n{'=' * 60}")
                    else:
                        full_text.append(text)

                    # Extract terminology
                    if extract_terminology:
                        self._extract_terms_from_text(text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                full_text.append(f"\n[ТАБЛИЦА {table_idx + 1}]")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))

            text_content = '\n'.join(full_text)

            if text_content:
                documents.append(
                    LangchainDocument(
                        page_content=text_content,
                        metadata={
                            "source": filename,
                            "type": self._categorize_document(filename),
                            "loaded_at": datetime.now().isoformat(),
                            "hash": self.get_file_hash(str(file_path))
                        }
                    )
                )

                # Update metadata
                self.loaded_docs[filename] = {
                    "loaded_at": datetime.now().isoformat(),
                    "hash": self.get_file_hash(str(file_path)),
                    "type": self._categorize_document(filename)
                }

        self._save_metadata()
        if extract_terminology:
            self._save_terminology()

        print(f"✅ Loaded {len(documents)} documents")
        return documents

    def _extract_terms_from_text(self, text: str):
        """Extract new insurance terminology from text."""
        # Simple extraction of key legal terms
        legal_markers = ["Статья", "Пункт", "Глава", "Раздел", "должны", "может быть"]
        for marker in legal_markers:
            if marker in text:
                # Store for later analysis
                pass

    def _categorize_document(self, filename: str) -> str:
        """Categorize document by filename."""
        categories = {
            "ГКРК": "Гражданский кодекс",
            "ОГПОВТС": "Обязательное страхование ТС",
            "перевозчика": "Страхование перевозчика",
            "туриста": "Страхование туристов",
            "экологии": "Страхование окружающей среды",
            "Нотариусы": "Страхование нотариусов",
            "опасные": "Страхование опасных объектов",
        }
        for keyword, category in categories.items():
            if keyword in filename:
                return category
        return "Прочие документы"

    def get_terminology_context(self) -> str:
        """Get terminology as context for LLM."""
        terms_list = []
        for term, definition in self.terminology.items():
            terms_list.append(f"- {term}: {definition}")
        return "\n".join(terms_list)

    def update_documents_from_adilet(self, adilet_urls: List[str]):
        """
        Placeholder for future integration with adilet.zan.kz
        Would download and parse documents from official KZ legislative database
        """
        print("⚠️ Adilet integration coming soon")
        pass

    def get_knowledge_base_info(self) -> Dict:
        """Get information about loaded knowledge base."""
        return {
            "total_documents": len(self.loaded_docs),
            "documents": self.loaded_docs,
            "terminology_count": len(self.terminology),
            "last_update": max(
                [doc.get("loaded_at") for doc in self.loaded_docs.values()]) if self.loaded_docs else None
        }
